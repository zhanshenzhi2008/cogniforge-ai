"""
DOCX document parser.
"""
import logging
from .base import BaseParser, ParsedDocument

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""

    SUPPORTED_EXTENSIONS = ['.docx']

    def __init__(self):
        self._docx_available = False
        try:
            import docx
            self.docx = docx
            self._docx_available = True
        except ImportError:
            logger.warning("python-docx not installed. DOCX parsing will be limited.")

    def parse(self, file_path: str) -> ParsedDocument:
        """Parse DOCX file and extract content."""
        try:
            text = self.extract_text(file_path)
            return ParsedDocument(
                content=text,
                metadata={"source": file_path, "type": "docx"},
                pages=0
            )
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise

    def extract_text(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        if not self._docx_available:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

        try:
            doc = self.docx.Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
